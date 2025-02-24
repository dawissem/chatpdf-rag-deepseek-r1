# # rag.py
# from langchain_core.globals import set_verbose, set_debug
# from langchain_ollama import ChatOllama, OllamaEmbeddings
# from langchain.schema.output_parser import StrOutputParser
# from langchain_community.vectorstores import Chroma
# from langchain_community.document_loaders import PyPDFLoader
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain.schema.runnable import RunnablePassthrough
# from langchain_community.vectorstores.utils import filter_complex_metadata
# from langchain_core.prompts import ChatPromptTemplate
# import logging

# set_debug(True)
# set_verbose(True)

# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)


# class ChatPDF:
#     """A class for handling PDF ingestion and question answering using RAG."""

#     def __init__(self, llm_model: str = "deepseek-r1:latest", embedding_model: str = "mxbai-embed-large"):
#         """
#         Initialize the ChatPDF instance with an LLM and embedding model.
#         """
#         self.model = ChatOllama(model=llm_model)
#         self.embeddings = OllamaEmbeddings(model=embedding_model)
#         self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=100)
#         self.prompt = ChatPromptTemplate.from_template(
#             """
#             You are a helpful assistant answering questions based on the uploaded document.
#             Context:
#             {context}
            
#             Question:
#             {question}
            
#             Answer concisely and accurately in three sentences or less.
#             """
#         )
#         self.vector_store = None
#         self.retriever = None

#     def ingest(self, pdf_file_path: str):
#         """
#         Ingest a PDF file, split its contents, and store the embeddings in the vector store.
#         """
#         logger.info(f"Starting ingestion for file: {pdf_file_path}")
#         docs = PyPDFLoader(file_path=pdf_file_path).load()
#         chunks = self.text_splitter.split_documents(docs)
#         chunks = filter_complex_metadata(chunks)

#         self.vector_store = Chroma.from_documents(
#             documents=chunks,
#             embedding=self.embeddings,
#             persist_directory="chroma_db",
#         )
#         logger.info("Ingestion completed. Document embeddings stored successfully.")

#     def ask(self, query: str, k: int = 5, score_threshold: float = 0.2):
#         """
#         Answer a query using the RAG pipeline.
#         """
#         if not self.vector_store:
#             raise ValueError("No vector store found. Please ingest a document first.")

#         if not self.retriever:
#             self.retriever = self.vector_store.as_retriever(
#                 search_type="similarity_score_threshold",
#                 search_kwargs={"k": k, "score_threshold": score_threshold},
#             )

#         logger.info(f"Retrieving context for query: {query}")
#         retrieved_docs = self.retriever.invoke(query)

#         if not retrieved_docs:
#             return "No relevant context found in the document to answer your question."

#         formatted_input = {
#             "context": "\n\n".join(doc.page_content for doc in retrieved_docs),
#             "question": query,
#         }

#         # Build the RAG chain
#         chain = (
#             RunnablePassthrough()  # Passes the input as-is
#             | self.prompt           # Formats the input for the LLM
#             | self.model            # Queries the LLM
#             | StrOutputParser()     # Parses the LLM's output
#         )

#         logger.info("Generating response using the LLM.")
#         return chain.invoke(formatted_input)

#     def clear(self):
#         """
#         Reset the vector store and retriever.
#         """
#         logger.info("Clearing vector store and retriever.")
#         self.vector_store = None
#         self.retriever = None


# first edition code 

























# from langchain_core.globals import set_verbose, set_debug
# from langchain_ollama import ChatOllama, OllamaEmbeddings
# from langchain.schema.output_parser import StrOutputParser
# from langchain_community.vectorstores import Chroma
# from langchain_community.document_loaders import PyPDFLoader
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain.schema.runnable import RunnablePassthrough
# from langchain_community.vectorstores.utils import filter_complex_metadata
# from langchain_core.prompts import ChatPromptTemplate
# import logging
# import pandas as pd
# import os
# import shutil
# from docx import Document

# set_debug(True)
# set_verbose(True)

# logging.basicConfig(level=logging.INFO)
# logger = logging.getLogger(__name__)


# class ChatDocument:
#     """A class for handling document ingestion and question answering using RAG."""

#     def __init__(self, llm_model: str = "deepseek-r1:latest", embedding_model: str = "mxbai-embed-large"):
#         """Initialize the ChatDocument instance with an LLM and embedding model."""
#         self.model = ChatOllama(model=llm_model)
#         self.embeddings = OllamaEmbeddings(model=embedding_model)
#         self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=4096, chunk_overlap=500)
#         self.prompt = ChatPromptTemplate.from_template(
#     """
#     Vous êtes un assistant qui doit extraire **toutes** les questions et leurs réponses du document fourni.

#     **IMPORTANT :**
#     - Il peut y avoir entre **15 et 20 questions**.
#     - Chaque question possède entre **1 et 5 réponses**.
#     - Ne générez que les questions réellement présentes dans le document.
#     - Copiez **exactement** les questions et leurs réponses sans reformulation.

#     **Format strict à respecter :**
    
#     1. **Question 1 ?**
#        - Réponse 1.1
#        - Réponse 1.2  
#        - Réponse 1.3 (si disponible)
#        - Réponse 1.4 (si disponible)
#        - Réponse 1.5 (si disponible)
#     2. **Question 2 ?**
#        - Réponse 2.1
#        - Réponse 2.2
#        - Réponse 2.3 (si disponible)
#        - Réponse 2.4 (si disponible)
#        - Réponse 2.5 (si disponible)
#     3. **Question 3 ?**
#        - Réponse 3.1
#        - Réponse 3.2
#        - Réponse 3.3 (si disponible)
#        - Réponse 3.4 (si disponible)
#        - Réponse 3.5 (si disponible)
#     ...
#     N. **Dernière question ?**
#        - Réponse N.1
#        - Réponse N.2 
#        - Réponse N.3 (si disponible)
#        - Réponse N.4 (si disponible)
#        - Réponse N.5 (si disponible)

#     **Instructions :**
#     - **Chaque question doit apparaître avec toutes ses réponses** (minimum 1, maximum 5).
#     - **Ne sautez aucune question ni aucune réponse**.
#     - **N’inventez pas de questions** si elles ne sont pas dans le document.
#     - **Respectez strictement le format ci-dessus** pour que les réponses soient bien organisées.

#     Contexte:
#     {context}

#     Question:
#     {question}

#     Respectez scrupuleusement le format ci-dessus et ne générez jamais plus de questions que celles réellement présentes dans le document.
#     """
# )

#         self.vector_store = None
#         self.retriever = None

#     def ingest(self, file_path: str):
#         """Ingest a document file (PDF, Word, or Excel)."""
#         file_extension = file_path.split(".")[-1].lower()
#         file_name = os.path.basename(file_path)  

#         logger.info(f"Checking if {file_name} is already in ChromaDB...")

#         if not self.vector_store:
#             self.vector_store = Chroma(persist_directory="chroma_db", embedding_function=self.embeddings)

#         existing_docs = self.vector_store.similarity_search(file_name, k=1)
#         if existing_docs:
#             logger.info(f"{file_name} is already in ChromaDB. Skipping ingestion.")
#             return  

#         logger.info(f"Ingesting {file_name}...")

#         if file_extension == "pdf":
#             docs = PyPDFLoader(file_path=file_path).load()
#             chunks = self.text_splitter.split_documents(docs)
#         elif file_extension == "docx":
#             text = self._extract_text_from_docx(file_path)
#             chunks = self.text_splitter.create_documents([text])
#         elif file_extension in ["xlsx", "csv"]:
#             df = pd.read_excel(file_path) if file_extension == "xlsx" else pd.read_csv(file_path)
#             text = df.to_string(index=False)
#             chunks = self.text_splitter.create_documents([text])
#         else:
#             raise ValueError("Unsupported file type. Please upload a PDF, Word, or Excel file.")

#         chunks = filter_complex_metadata(chunks)
#         self.vector_store.add_documents(chunks)  

#         logger.info(f"{file_name} has been successfully ingested into ChromaDB.")

#     def ingest_text(self, text: str):
#         """Ingest plain text (from Word or Excel files)."""
#         if not text.strip():
#             logger.warning("Empty text received, skipping ingestion.")
#             return

#         chunks = self.text_splitter.create_documents([text])
#         chunks = filter_complex_metadata(chunks)

#         if not self.vector_store:
#             self.vector_store = Chroma(persist_directory="chroma_db", embedding_function=self.embeddings)

#         self.vector_store.add_documents(chunks)
#         logger.info("Text ingestion completed. Document embeddings stored successfully.")

#     def clear(self):
#         """Reset the vector store and retriever."""
#         self.vector_store = None
#         self.retriever = None
#         logger.info("ChatDocument has been cleared.")

#     def clear_chroma_db(self):
#         """Clears all embeddings from ChromaDB without deleting the database structure."""
#         if self.vector_store:
#             self.vector_store.delete_collection()
#             logger.info("ChromaDB embeddings have been deleted but the database remains intact.")
#         else:
#             logger.info("No vector store found. Nothing to delete.")

#     def ask(self, query: str, k: int = 5, score_threshold: float = 0.2):
#         """Answer a query using the RAG pipeline."""
#         if not self.vector_store:
#             raise ValueError("No vector store found. Please ingest a document first.")

#         if not self.retriever:
#             self.retriever = self.vector_store.as_retriever(
#                 search_type="similarity_score_threshold",
#                 search_kwargs={"k": k, "score_threshold": score_threshold},
#             )

#         logger.info(f"Retrieving context for query: {query}")
#         retrieved_docs = self.retriever.invoke(query)

#         if not retrieved_docs:
#             return "No relevant context found in the document to answer your question."

#         formatted_input = {
#             "context": "\n\n".join(doc.page_content for doc in retrieved_docs),
#             "question": query,
#         }

#         chain = (
#             RunnablePassthrough()  
#             | self.prompt          
#             | self.model          
#             | StrOutputParser()    
#         )

#         logger.info("Generating response using the LLM.")
#         return chain.invoke(formatted_input)

#     def _extract_text_from_docx(self, file_path: str):
#         """Extract text from a Word (.docx) file."""
#         doc = Document(file_path)
#         return "\n".join([para.text for para in doc.paragraphs])



#3rd version 
from langchain_core.globals import set_verbose, set_debug
from langchain_ollama import ChatOllama, OllamaEmbeddings
from langchain.schema.output_parser import StrOutputParser
from langchain_community.vectorstores import Chroma
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema.runnable import RunnablePassthrough
from langchain_community.vectorstores.utils import filter_complex_metadata
from langchain_core.prompts import ChatPromptTemplate
import logging
import pandas as pd
import os
from docx import Document

set_debug(True)
set_verbose(True)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ChatDocument:
    """A class for handling document ingestion and question answering using RAG."""

    def __init__(self, llm_model: str = "deepseek-r1:latest", embedding_model: str = "mxbai-embed-large"):
        """Initialize the ChatDocument instance with an LLM and embedding model."""
        self.model = ChatOllama(model=llm_model)
        self.embeddings = OllamaEmbeddings(model=embedding_model)
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=4096, chunk_overlap=500)
        self.prompt = ChatPromptTemplate.from_template(
            """
            Vous êtes un assistant qui doit répondre **exclusivement en français** et qui doit extraire **toutes** les questions et leurs réponses du document fourni.

            **IMPORTANT :**
            - Les questions sont **en gras** dans le document et les propositions de réponses sont en dessous de chaque question.
            - **Ne générez jamais de nouvelles questions ou réponses.** Utilisez uniquement celles présentes dans le document.
            - **Respectez strictement le texte des questions et des réponses.** Ne reformulez rien.
            - **Si une question a au maximum 5 proposition de réponses, affichez uniquement celles disponibles.**
            - **Affichez uniquement le contenu pertinent sous forme de liste.** Ne fournissez aucune explication, aucun commentaire, et aucune interprétation.
            - **Évitez d'ajouter des symboles inutiles comme `??` '!!' ou `think>`**.

            **Format strict à respecter :**
                
                1. **Question 1 :** (Texte exact de la question 1 en gras)
                    - **Réponse 1.1** : (Texte exact de la réponse 1.1)
                    - **Réponse 1.2** : (Texte exact de la réponse 1.2) (si disponible)
                    - **Réponse 1.3** : (Texte exact de la réponse 1.3) (si disponible)
                    - **Réponse 1.4** : (Texte exact de la réponse 1.4) (si disponible)
                    - **Réponse 1.5** : (Texte exact de la réponse 1.5) (si disponible)

                2. **Question 2 :** ((Texte exact de la question 2 en gras)
                    - **Réponse 2.1** : (Texte exact de la réponse 2.1)
                    - **Réponse 2.2** : (Texte exact de la réponse 2.2) (si disponible)
                    - **Réponse 2.3** : (Texte exact de la réponse 2.3) (si disponible)
                    - **Réponse 2.4** : (Texte exact de la réponse 2.4) (si disponible)
                    - **Réponse 2.5** : (Texte exact de la réponse 2.5) (si disponible)
                ...
                N. **Question N :** (Texte exact de la question N en gras)
                    - **Réponse N.1** : (Texte exact de la réponse N.1)
                    - **Réponse N.2** : (Texte exact de la réponse N.2) (si disponible)
                    - **Réponse N.3** : (Texte exact de la réponse N.3) (si disponible)
                    - **Réponse N.4** : (Texte exact de la réponse N.4) (si disponible)
                    - **Réponse N.5** : (Texte exact de la réponse N.5) (si disponible)

                **Contexte fourni par le document :**
                {context}

                **Question de l'utilisateur :**
                {question}

            **Règles supplémentaires :**
                - **Ne sautez aucune question ni aucune proposition de réponse**. Chaque question avec ses réponses doit être affichée.
                - Assurez-vous que **le format et l'ordre des questions** soient respectés tels qu'ils apparaissent dans le document.
                - **Si aucune réponse n'est présente pour une question**, **affichez simplement la question sans réponse** (par exemple, "Réponse 1.1 : Non disponible").

                Respectez scrupuleusement le format ci-dessus et **ne générez jamais plus de questions** que celles réellement présentes dans le document.            """
)


        self.vector_store = None
        self.retriever = None

    def ingest(self, file_path: str):
        """Ingest a document file (PDF, Word, or Excel)."""
        file_extension = file_path.split(".")[-1].lower()
        file_name = os.path.basename(file_path)  

        logger.info(f"Checking if {file_name} is already in ChromaDB...")

        if not self.vector_store:
            self.vector_store = Chroma(persist_directory="chroma_db", embedding_function=self.embeddings)

        existing_docs = self.vector_store.similarity_search(file_name, k=1)
        if existing_docs:
            logger.info(f"{file_name} is already in ChromaDB. Skipping ingestion.")
            return  

        logger.info(f"Ingesting {file_name}...")

        if file_extension == "pdf":
            docs = PyPDFLoader(file_path=file_path).load()
            chunks = self.text_splitter.split_documents(docs)
        elif file_extension == "docx":
            text = self._extract_text_from_docx(file_path)
            self.ingest_text(text)
        elif file_extension in ["xlsx", "csv"]:
            df = pd.read_excel(file_path) if file_extension == "xlsx" else pd.read_csv(file_path)
            text = df.to_string(index=False)
            self.ingest_text(text)
        else:
            raise ValueError("Unsupported file type. Please upload a PDF, Word, or Excel file.")

        logger.info(f"{file_name} has been successfully ingested into ChromaDB.")

    def ingest_text(self, text: str):
        """Ingest plain text (from Word or Excel files) into ChromaDB."""
        if not text.strip():
            logger.warning("Empty text received, skipping ingestion.")
            return

        chunks = self.text_splitter.create_documents([text])
        chunks = filter_complex_metadata(chunks)

        if not self.vector_store:
            self.vector_store = Chroma(persist_directory="chroma_db", embedding_function=self.embeddings)

        self.vector_store.add_documents(chunks)
        logger.info("Text ingestion completed. Document embeddings stored successfully.")

    def ask(self, query: str, k: int = 5, score_threshold: float = 0.2):
        """Retrieve all questions and their actual answers from the document without making assumptions."""

        if not self.vector_store:
            raise ValueError("Aucune base de données vectorielle trouvée. Veuillez d'abord ingérer un document.")

        if not self.retriever:
            self.retriever = self.vector_store.as_retriever(
                search_type="similarity_score_threshold",
                search_kwargs={"k": k, "score_threshold": score_threshold},
            )

        logger.info(f"Récupération du contexte pour la question : {query}")
        retrieved_docs = self.retriever.invoke(query)

        if not retrieved_docs:
            return "Aucun contexte pertinent trouvé dans le document pour répondre à votre question."

        formatted_input = {
            "context": "\n\n".join(doc.page_content for doc in retrieved_docs),
            "question": query,
        }

        chain = (
            RunnablePassthrough()
            | self.prompt
            | self.model
            | StrOutputParser()
        )

        logger.info("Génération de la réponse en français.")

        response = chain.invoke(formatted_input)

        # Vérifier si la réponse contient réellement des questions et réponses extraites
        if "Question" not in response or "Réponse" not in response:
            return "Erreur : Les questions et réponses n'ont pas été correctement extraites du document. Veuillez vérifier le contenu du fichier."

        return response.strip()


    def clear(self):
        """Clears the vector store and retriever to reset the document processing."""
        self.vector_store = None
        self.retriever = None
        logger.info("ChatDocument has been cleared.")

    def clear_chroma_db(self):
        """Clears all embeddings from ChromaDB without deleting the database structure."""
        if self.vector_store:
            self.vector_store.delete_collection()
            logger.info("ChromaDB embeddings have been deleted but the database remains intact.")
        else:
            logger.info("No vector store found. Nothing to delete.")

    def _extract_text_from_docx(self, file_path: str):
        """Extract text from a Word (.docx) file."""
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
