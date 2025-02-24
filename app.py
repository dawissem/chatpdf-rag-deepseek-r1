# # app.py
# import os
# import tempfile
# import time
# import streamlit as st
# from streamlit_chat import message
# from rag import ChatPDF
 

# st.set_page_config(page_title="RAG with Local DeepSeek R1")


# def display_messages():
#     """Display the chat history."""
#     st.subheader("Chat History")
#     for i, (msg, is_user) in enumerate(st.session_state["messages"]):
#         message(msg, is_user=is_user, key=str(i))
#     st.session_state["thinking_spinner"] = st.empty()


# def process_input():
#     """Process the user input and generate an assistant response."""
#     if st.session_state["user_input"] and len(st.session_state["user_input"].strip()) > 0:
#         user_text = st.session_state["user_input"].strip()
#         with st.session_state["thinking_spinner"], st.spinner("Thinking..."):
#             try:
#                 agent_text = st.session_state["assistant"].ask(
#                     user_text,
#                     k=st.session_state["retrieval_k"],
#                     score_threshold=st.session_state["retrieval_threshold"],
#                 )
#             except ValueError as e:
#                 agent_text = str(e)

#         st.session_state["messages"].append((user_text, True))
#         st.session_state["messages"].append((agent_text, False))


# def read_and_save_file():
#     """Handle file upload and ingestion."""
#     st.session_state["assistant"].clear()
#     st.session_state["messages"] = []
#     st.session_state["user_input"] = ""

#     for file in st.session_state["file_uploader"]:
#         with tempfile.NamedTemporaryFile(delete=False) as tf:
#             tf.write(file.getbuffer())
#             file_path = tf.name

#         with st.session_state["ingestion_spinner"], st.spinner(f"Ingesting {file.name}..."):
#             t0 = time.time()
#             st.session_state["assistant"].ingest(file_path)
#             t1 = time.time()

#         st.session_state["messages"].append(
#             (f"Ingested {file.name} in {t1 - t0:.2f} seconds", False)
#         )
#         os.remove(file_path)


# def page():
#     """Main app page layout."""
#     if len(st.session_state) == 0:
#         st.session_state["messages"] = []
#         st.session_state["assistant"] = ChatPDF()

#     st.header("RAG with Local DeepSeek R1")

#     st.subheader("Upload a Document")
#     st.file_uploader(
#         "Upload a PDF document",
#         type=["pdf"],
#         key="file_uploader",
#         on_change=read_and_save_file,
#         label_visibility="collapsed",
#         accept_multiple_files=True,
#     )

#     st.session_state["ingestion_spinner"] = st.empty()

#     # Retrieval settings
#     st.subheader("Settings")
#     st.session_state["retrieval_k"] = st.slider(
#         "Number of Retrieved Results (k)", min_value=1, max_value=10, value=5
#     )
#     st.session_state["retrieval_threshold"] = st.slider(
#         "Similarity Score Threshold", min_value=0.0, max_value=1.0, value=0.2, step=0.05
#     )

#     # Display messages and text input
#     display_messages()
#     st.text_input("Message", key="user_input", on_change=process_input)

#     # Clear chat
#     if st.button("Clear Chat"):
#         st.session_state["messages"] = []
#         st.session_state["assistant"].clear()


# if __name__ == "__main__":
#     page()
import os
import tempfile
import time
import streamlit as st
import pandas as pd
from docx import Document  # Library for reading Word files (.docx)
from streamlit_chat import message
from rag import ChatDocument  # Updated to support multiple file types

st.set_page_config(page_title="RAG with Local DeepSeek R1")


def display_messages():
    """Display the chat history."""
    st.subheader("Chat History")
    for i, (msg, is_user) in enumerate(st.session_state["messages"]):
        message(msg, is_user=is_user, key=str(i))
    st.session_state["thinking_spinner"] = st.empty()


def process_input():
    """Process the user input and generate an assistant response."""
    if st.session_state["user_input"] and len(st.session_state["user_input"].strip()) > 0:
        user_text = st.session_state["user_input"].strip()
        with st.session_state["thinking_spinner"], st.spinner("Thinking..."):
            try:
                agent_text = st.session_state["assistant"].ask(
                    user_text,
                    k=st.session_state["retrieval_k"],
                    score_threshold=st.session_state["retrieval_threshold"],
                )
            except ValueError as e:
                agent_text = str(e)

        st.session_state["messages"].append((user_text, True))
        st.session_state["messages"].append((agent_text, False))


def extract_text_from_docx(file_path):
    """Extract text from a Word (.docx) file."""
    doc = Document(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    return full_text


def read_and_save_file():
    """Handle file upload and ingestion for PDFs, Excel, and Word files."""
    st.session_state["assistant"].clear()
    st.session_state["messages"] = []
    st.session_state["user_input"] = ""

    for file in st.session_state["file_uploader"]:
        file_extension = file.name.split(".")[-1].lower()

        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as tf:
            tf.write(file.getbuffer())
            file_path = tf.name

        with st.session_state["ingestion_spinner"], st.spinner(f"Ingesting {file.name}..."):
            t0 = time.time()

            if file_extension == "pdf":
                st.session_state["assistant"].ingest(file_path)  # Process PDF
            elif file_extension == "docx":
                doc_text = extract_text_from_docx(file_path)
                st.session_state["assistant"].ingest_text(doc_text)  # Process Word Document
            elif file_extension in ["xlsx", "csv"]:
                df = pd.read_excel(file_path) if file_extension == "xlsx" else pd.read_csv(file_path)
                text_data = df.to_string(index=False)
                st.session_state["assistant"].ingest_text(text_data)  # Process Excel as text

            t1 = time.time()

        st.session_state["messages"].append(
            (f"Ingested {file.name} in {t1 - t0:.2f} seconds", False)
        )
        os.remove(file_path)

def page():
    """Main app page layout."""
    if len(st.session_state) == 0:
        st.session_state["messages"] = []
        st.session_state["assistant"] = ChatDocument()  # Updated class to handle PDFs, Excel, and Word files

    st.header("RAG with Local DeepSeek R1")

    st.subheader("Upload a Document")
    st.file_uploader(
        "Upload a PDF, Word, or Excel document",
        type=["pdf", "docx", "xlsx", "csv"],
        key="file_uploader",
        on_change=read_and_save_file,
        label_visibility="collapsed",
        accept_multiple_files=True,
    )

    st.session_state["ingestion_spinner"] = st.empty()

    # Retrieval settings
    st.subheader("Settings")
    st.session_state["retrieval_k"] = st.slider(
        "Number of Retrieved Results (k)", min_value=1, max_value=10, value=5
    )
    st.session_state["retrieval_threshold"] = st.slider(
        "Similarity Score Threshold", min_value=0.0, max_value=1.0, value=0.2, step=0.05
    )

    # Display messages and text input
    display_messages()

    st.text_input("Message", key="user_input", on_change=process_input)

    # Clear chat
    if st.button("Clear Chat"):
        st.session_state["messages"] = []
        st.session_state["assistant"].clear()

    if st.button("Clear ChromaDB"):
        st.session_state["assistant"].clear_chroma_db()
        st.success("ChromaDB has been cleared successfully!")


# st.session_state["assistant"].clear_chroma_db()  # Clears database when the app starts

if __name__ == "__main__":
    page()
